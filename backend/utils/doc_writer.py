import os
import datetime
from docx import Document
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

def generate_word_document(grouped_articles: List[Dict[str, Any]]) -> str:
    """
    Generate a Word document with the social media content plan.
    
    Args:
        grouped_articles: List of feed groups with articles and generated content
        
    Returns:
        Filename of the generated document
    """
    try:
        # Generate filename with timestamp
        now = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        doc_filename = f"social_plan_{now}.docx"
        
        # Create documents directory if it doesn't exist
        documents_dir = os.path.join(os.getcwd(), "documents")
        os.makedirs(documents_dir, exist_ok=True)
        doc_path = os.path.join(documents_dir, doc_filename)
        
        # Create Word document
        doc = Document()
        doc.add_heading("Dental Clinic Social Media Content Plan", 0)
        doc.add_paragraph(f"Generated: {now}")
        
        # Add content for each feed
        for feed in grouped_articles:
            doc.add_heading(feed['feed_title'], level=1)
            doc.add_paragraph(f"Feed URL: {feed['feed_url']}")
            
            for article in feed['articles']:
                doc.add_paragraph(f"- {article['title']} ({article['link']})")
                
                if article.get('summary'):
                    doc.add_paragraph(f"  Summary: {article['summary']}")
                    doc.add_paragraph(f"  IG Caption: {article['caption']}")
                    doc.add_paragraph(f"  Image Prompt: {article['imagePrompt']}")
        
        # Save document
        doc.save(doc_path)
        logger.info(f"Successfully generated Word document: {doc_filename}")
        return doc_filename
        
    except Exception as e:
        logger.error(f"Failed to generate Word document: {str(e)}")
        raise 