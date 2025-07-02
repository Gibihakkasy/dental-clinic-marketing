from urllib.parse import urlparse, parse_qs, urlencode, urlunparse

def clean_knowyourteeth_link(link: str) -> str:
    parsed = urlparse(link)
    if "knowyourteeth.com" in parsed.netloc:
        query = parse_qs(parsed.query)
        cleaned = {}
        # Preserve relevant parameters if present
        for key in ("abc", "iid", "aid"):
            if key in query:
                cleaned[key] = query[key][0]
        new_query = urlencode(cleaned)
        return urlunparse((
            "https",
            "knowyourteeth.com",
            parsed.path,
            "",
            new_query,
            ""
        ))
    return link
from dotenv import load_dotenv
load_dotenv()

import os
import json
import logging
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import requests
import openai
from bs4 import BeautifulSoup
import feedparser
import datetime
from docx import Document
from fastapi.responses import FileResponse

# Dental news fetching utilities (inlined from fetch_news.py)
import feedparser
import datetime

# Load API key
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Model constants
SUMMARIZE_MODEL = "gpt-4o"
CAPTION_MODEL = "gpt-4.1-nano"
IMAGE_MODEL = "gpt-4.1-nano"

# Social Media Planner agent config
AGENT = {
    "name": "Social Media Planner",
    "provider": "openai",
    "model": "gpt-4o"
}

# RSS feeds
RSS_FEEDS = [
    "http://www.agd.org/myagd/subscriptions/rss/kyt_hottopics.xml",
    "http://www.agd.org/myagd/subscriptions/rss/kyt_factoidweek.xml",
    "https://www.dentalhealth.org/handlers/rss.ashx?feed=1",
    "https://askthedentist.com/feed/"
]

def extract_article_text(url: str, max_paragraphs: int = 5) -> str:
    try:
        resp = requests.get(url, timeout=5)
        soup = BeautifulSoup(resp.text, "html.parser")
        paragraphs = soup.find_all("p")
        return " ".join(p.get_text().strip() for p in paragraphs[:max_paragraphs])
    except Exception:
        return ""

def fetch_unique_rss_news():
    entries = []
    seen_titles = set()
    seen_links = set()
    for url in RSS_FEEDS:
        feed = feedparser.parse(url)
        for entry in feed.entries:
            title = entry.get("title")
            raw_link = entry.get("link")
            if isinstance(raw_link, list):
                raw_link = raw_link[0] if raw_link else None
            # Only clean if raw_link is a string
            if not isinstance(raw_link, str):
                continue
            link = clean_knowyourteeth_link(raw_link)
            if not title or not link:
                continue
            if title in seen_titles or link in seen_links:
                continue
            seen_titles.add(title)
            seen_links.add(link)
            published = None
            if hasattr(entry, 'published_parsed') and entry.published_parsed and isinstance(entry.published_parsed, tuple):
                try:
                    published = datetime.datetime(*entry.published_parsed[:6])
                except Exception:
                    published = datetime.datetime.now()
            else:
                published = datetime.datetime.now()
            entries.append({
                "title": title,
                "link": link,
                "published": published
            })
    sorted_entries = sorted(entries, key=lambda x: x["published"], reverse=True)
    latest = sorted_entries[:5]
    return [{"title": e["title"], "link": e["link"]} for e in latest]

# New: fetch grouped articles (3 per feed)
def fetch_grouped_rss_news():
    grouped = []
    for url in RSS_FEEDS:
        feed = feedparser.parse(url)
        feed_title = url
        if hasattr(feed, 'feed') and isinstance(feed.feed, dict):
            feed_title = feed.feed.get("title", url)
        articles = []
        seen_titles = set()
        seen_links = set()
        for entry in feed.entries:
            if len(articles) >= 5:
                break
            title = entry.get("title")
            raw_link = entry.get("link")
            if isinstance(raw_link, list):
                raw_link = raw_link[0] if raw_link else None
            if not isinstance(raw_link, str):
                continue
            link = clean_knowyourteeth_link(raw_link)
            if not title or not link:
                continue
            if title in seen_titles or link in seen_links:
                continue
            seen_titles.add(title)
            seen_links.add(link)
            published = None
            if hasattr(entry, 'published_parsed') and entry.published_parsed and isinstance(entry.published_parsed, tuple):
                try:
                    published = datetime.datetime(*entry.published_parsed[:6])
                except Exception:
                    published = datetime.datetime.now()
            else:
                published = datetime.datetime.now()
            articles.append({
                "title": title,
                "link": link,
                "published": published
            })
        grouped.append({
            "feed_title": feed_title,
            "feed_url": url,
            "articles": articles
        })
    return grouped

# FastAPI app setup
app = FastAPI()

origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class SummarizeRequest(BaseModel):
    user_message: Optional[str] = None

@app.get("/bots")
def get_bots():
    return [{"name": AGENT["name"], "role": "Creates Instagram content plan from latest dental news."}]

@app.post("/generate_social_plan")
async def generate_social_plan(request: dict):
    try:
        if not OPENAI_API_KEY:
            raise Exception("OPENAI_API_KEY not set")
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        # Expect request to have 'selected' as a list of {feed_url, article_link}
        selected = request.get('selected', [])
        grouped = fetch_grouped_rss_news()
        # Build a lookup for selected articles
        selected_lookup = {(s['feed_url'], s['article_link']) for s in selected}
        # Prepare results for frontend (optional)
        results = []
        for feed in grouped:
            for article in feed['articles']:
                is_selected = (feed['feed_url'], article['link']) in selected_lookup
                if is_selected:
                    article_text = extract_article_text(article["link"])
                    summary_prompt = (
                        f"Summarize the following dental news article for dentists audience in Bahasa Indonesia.\n"
                        f"Title: {article['title']}\n"
                        f"URL: {article['link']}\n"
                        f"Return a comprehensive summary. Without flavor text, e.g. here's the summary:, or repeating the title in the summary."
                    )
                    # Use OpenAI Response API with web_search tool for detailed summarization
                    summary_resp = client.responses.create(
                        model=SUMMARIZE_MODEL,
                        input=summary_prompt,
                        tools=[{"type": "web_search"}]
                    )
                    summary = getattr(summary_resp, 'output_text', None)
                    if summary:
                        summary = summary.strip()
                    else:
                        summary = "(No summary returned)"
                    caption_prompt = (
                        f"Given this summary of a dental news article, write a highly engaging Instagram caption for a dental clinic's patients in Bahasa Indonesia. Don't use em-dashes. Always include call to action. \n"
                        f"Summary: {summary}\n"
                    )
                    caption_resp = client.chat.completions.create(
                        model=CAPTION_MODEL,
                        messages=[{"role": "user", "content": caption_prompt}]
                    )
                    caption = getattr(caption_resp.choices[0].message, 'content', None)
                    if caption:
                        caption = caption.strip()
                    else:
                        caption = "(No caption returned)"
                    image_prompt_prompt = (
                        f"Given this summary, write a prompt for an image generation AI to create an image for a dental clinic instagram post. The audience is potential patient of an Indonesian dental clinic. Prioritize photorealistic image of Indonesian. 2D cartoon can be added as flare or decoration. Return only the prompt, no other text.\n"
                        f"Summary: {summary}\n"
                    )
                    image_prompt_resp = client.chat.completions.create(
                        model=IMAGE_MODEL,
                        messages=[{"role": "user", "content": image_prompt_prompt}]
                    )
                    image_prompt = getattr(image_prompt_resp.choices[0].message, 'content', None)
                    if image_prompt:
                        image_prompt = image_prompt.strip()
                    else:
                        image_prompt = "(No image prompt returned)"
                    article['summary'] = summary
                    article['caption'] = caption
                    article['imagePrompt'] = image_prompt
                else:
                    article['summary'] = None
                    article['caption'] = None
                    article['imagePrompt'] = None
        # Generate Word document
        now = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        doc_filename = f"social_plan_{now}.docx"
        documents_dir = os.path.join(os.getcwd(), "documents")
        os.makedirs(documents_dir, exist_ok=True)
        doc_path = os.path.join(documents_dir, doc_filename)
        doc = Document()
        doc.add_heading("Dental Clinic Social Media Content Plan", 0)
        doc.add_paragraph(f"Generated: {now}")
        for feed in grouped:
            doc.add_heading(feed['feed_title'], level=1)
            doc.add_paragraph(f"Feed URL: {feed['feed_url']}")
            for article in feed['articles']:
                doc.add_paragraph(f"- {article['title']} ({article['link']})")
                if article['summary']:
                    doc.add_paragraph(f"  Summary: {article['summary']}")
                    doc.add_paragraph(f"  IG Caption: {article['caption']}")
                    doc.add_paragraph(f"  Image Prompt: {article['imagePrompt']}")
        doc.save(doc_path)
        return {"file": doc_filename, "grouped": grouped}
    except Exception as e:
        logging.error(f"Error generating social plan: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.get("/download/{filename}")
def download_file(filename: str):
    file_path = os.path.join("backend", filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)

@app.get("/get_rss_articles")
def get_rss_articles():
    return fetch_unique_rss_news()

@app.get("/get_rss_articles_grouped")
def get_rss_articles_grouped():
    return fetch_grouped_rss_news()
